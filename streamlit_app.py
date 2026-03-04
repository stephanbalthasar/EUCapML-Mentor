# streamlit_app.py
# Minimal UI to exercise both engines using the booklet index from the private repo.

import streamlit as st

# --- Load booklet index (server-side; users never see this file) ---
from app.bootstrap_booklet import load_booklet_index
INDEX = load_booklet_index()  # {"paragraphs": [...], "chapters": [...]}

from app.bootstrap_cases import load_cases
CASES = load_cases()

from mentor.booklet.retriever import ParagraphRetriever, ChapterRetriever
from mentor.engines.chat_engine import ChatEngine
from mentor.engines.feedback_engine import FeedbackEngine
from mentor.llm.groq import GroqClient

st.set_page_config(page_title="EUCapML Mentor", page_icon="⚖️", layout="wide")

# Optional tiny banner (remove once you’re confident)
st.caption(f"📖 Booklet loaded — {len(INDEX['chapters'])} chapters, {len(INDEX['paragraphs'])} numbered paragraphs.")

# --- Build retrievers once ---
para_retriever = ParagraphRetriever(INDEX["paragraphs"])
chap_retriever = ChapterRetriever(INDEX["chapters"])

# --- LLM client ---
llm_api_key = st.secrets.get("GROQ_API_KEY")
if not llm_api_key:
    st.error("Missing GROQ_API_KEY in secrets.")
    st.stop()
llm = GroqClient(api_key=llm_api_key)

# --- Engines ---
chat_engine = ChatEngine(
    llm=llm,
    booklet_index=INDEX,
    booklet_retriever=chap_retriever,  # ChatEngine uses chapter-level grounding
    web_retriever=None                 # add later if you want web RAG
)
feedback_engine = FeedbackEngine(llm=llm)

# --- Sidebar controls ---
with st.sidebar:
    model = st.selectbox("Model", ["llama-3.1-8b-instant", "llama-3.3-70b-versatile"], index=0)
    temp  = st.slider("Temperature", 0.0, 1.0, 0.2, 0.05)
    if st.button("Reload booklet index (server cache)"):
        st.cache_data.clear()
        st.success("Re-loaded. Re-run the action to use the latest JSON.")

# --- Tabs: Feedback + Tutor chat ---
tab_feedback, tab_chat = st.tabs(["📝 Feedback", "💬 Tutor chat"])

# Small helper: persist latest run per case+question
def _key(case_id: str, q_label: str) -> str:
    return f"{case_id}::{q_label}"

# --- REPLACEMENT FEEDBACK TAB (full block) ----

with tab_feedback:

    st.subheader("Exam Assistant")

    # -----------------------------
    # 1. CASE SELECTION (kept unchanged)
    # -----------------------------
    case_titles = [c.get("title", c.get("id", "Untitled case")) for c in CASES]
    sel_case_title = st.selectbox("Select exam case", case_titles, index=0)
    sel_case = next(c for c in CASES if c.get("title", c.get("id")) == sel_case_title)
    sel_case_id = sel_case.get("id", "unknown")

    q_count = int(sel_case.get("question_count", 1))
    q_labels = [f"Question {i+1}" for i in range(max(1, q_count))]
    q_label = st.selectbox("Which question are you working on?", q_labels, index=0)
    q_index = q_labels.index(q_label)

    # -----------------------------
    # 2. CASE DESCRIPTION
    # -----------------------------
    st.markdown("### Case description")
    st.write(sel_case.get("description", "—"))

    st.divider()

    # -----------------------------
    # 3. WORKFLOW CHOICE
    # -----------------------------
    workflow = st.radio(
        "Choose your workflow:",
        ["Help me prepare an answer", "I have an answer ready to submit"],
        horizontal=False
    )

    st.divider()


    # -----------------------------
    # 4. PLAN WORKFLOW
    # -----------------------------
    if workflow == "Help me prepare an answer":
        st.markdown("## Plan your answer")
        st.markdown("The app will help you build a structured outline based on the case and model solution.")

        # Load model answer slice
        sections = sel_case.get("model_answer_sections") or []
        model_slice = sections[q_index] if (0 <= q_index < len(sections)) else ""

        if st.button("Generate plan", type="primary"):
            with st.spinner("Thinking..."):
                plan = feedback_engine.plan_answer(
                    case_text=sel_case.get("description", ""),
                    question=q_label,
                    model_answer_slice=model_slice,
                    booklet_text="",       # your design: no booklet grounding in plan
                    model=model,
                    temperature=temp
                )
            st.session_state["plan_output"] = plan

        # Display generated plan
        if "plan_output" in st.session_state:
            st.markdown("### Suggested solution structure")
            st.markdown(st.session_state["plan_output"])


    # -----------------------------
    # 5. EVALUATE WORKFLOW
    # -----------------------------
    if workflow == "I have an answer ready to submit":

        st.markdown("## Submit your exam answer")

        # text input
        answer = st.text_area(
            "Your answer",
            height=240,
            key=f"answer::{sel_case_id}::{q_label}"
        )

        # on evaluate
        if st.button("Evaluate my answer", type="primary"):
            sections = sel_case.get("model_answer_sections") or []
            auto_slice = sections[q_index] if (0 <= q_index < len(sections)) else None
            effective_model = (auto_slice or "").strip()

            if not effective_model or not answer.strip():
                st.warning("Missing model answer slice or student answer.")
            else:
                with st.spinner("Evaluating..."):
                    fb = feedback_engine.evaluate_answer(
                        student_answer=answer,
                        model_answer=effective_model,
                        model=model,
                        temperature=temp
                    )

                # persist all relevant state
                st.session_state["exam_answer"] = answer
                st.session_state["exam_feedback"] = fb
                st.session_state["chat_history"] = []    # reset chat thread

        # SHOW RESULTS AFTER EVALUATION
        if "exam_feedback" in st.session_state:

            st.markdown("## Your submitted answer")
            st.markdown(st.session_state["exam_answer"])

            st.markdown("## Structured feedback")
            st.markdown(st.session_state["exam_feedback"])

            # -----------------------------
            # DOCX DOWNLOAD
            # -----------------------------
            from docx import Document
            from io import BytesIO

            def make_docx():
                doc = Document()
                doc.add_heading(f"Feedback – {sel_case_title} – {q_label}", level=1)

                doc.add_heading("Student Answer", level=2)
                doc.add_paragraph(st.session_state["exam_answer"])

                doc.add_heading("Feedback", level=2)
                doc.add_paragraph(st.session_state["exam_feedback"])

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf

            st.download_button(
                "📄 Download feedback (.docx)",
                data=make_docx(),
                file_name=f"feedback_{sel_case_id}_{q_label.replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.divider()

            # -----------------------------
            # 6. FOLLOW-UP CHAT
            # -----------------------------
            st.markdown("## Follow-up discussion")

            # show history
            for role, msg in st.session_state["chat_history"]:
                if role == "student":
                    st.markdown(f"**You:** {msg}")
                else:
                    st.markdown(f"**Tutor:** {msg}")

            follow_q = st.text_area("Your follow-up question", height=120)

            if st.button("Send follow-up"):
                if follow_q.strip():

                    # add user's message
                    st.session_state["chat_history"].append(("student", follow_q))

                    # build conversation context
                    context = {
                        "student_answer": st.session_state["exam_answer"],
                        "feedback": st.session_state["exam_feedback"],
                        "history": st.session_state["chat_history"],
                    }

                    # new engine call
                    reply = feedback_engine.follow_up_with_history(
                        question=follow_q,
                        context=context,
                        model=model,
                        temperature=temp
                    )

                    # store bot reply
                    st.session_state["chat_history"].append(("tutor", reply))

                    # force re-render
                    st.rerun()

# --- Tutor chat (separate, uncluttered) ---
with tab_chat:
    st.subheader("Tutor chat (booklet‑grounded)")
    q = st.text_area("Your question", height=140, placeholder="e.g., What is 'inside information' under MAR?")
    if st.button("Ask", key="chat_btn"):
        if not q.strip():
            st.warning("Please enter a question.")
        else:
            with st.spinner("Thinking..."):
                reply = chat_engine.answer(q, model=model, temperature=temp, max_tokens=800)
            st.markdown(reply)
