# scripts/build_booklet_index.py
import json
import sys
from pathlib import Path
from docx import Document

def build_index(docx_path: str) -> dict:
    doc = Document(docx_path)

    # 1) Chapters from Heading 1 (support EN+DE style names)
    chapters = []
    current_title, current_num, buf = None, 0, []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style in ("Heading 1", "Überschrift 1"):
            if current_title is not None:
                chapters.append({
                    "chapter_num": current_num,
                    "title": current_title,
                    "text": "\n".join(buf).strip(),
                })
            current_num += 1
            current_title = text or f"Chapter {current_num}"
            buf = []
        else:
            if current_title is not None and text:
                buf.append(text)
    if current_title is not None:
        chapters.append({
            "chapter_num": current_num,
            "title": current_title,
            "text": "\n".join(buf).strip(),
        })

    # 2) Paragraphs from "Standard with para numbering" (sequential 1..N)
    paragraphs, para_counter, chapter_cursor = [], 0, 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if style in ("Heading 1", "Überschrift 1"):
            chapter_cursor += 1
            continue
        if style == "Standard with para numbering" and text:
            para_counter += 1
            paragraphs.append({
                "para_num": para_counter,
                "text": text,
                "chapter_num": chapter_cursor or None,
                "chapter_title": chapters[chapter_cursor - 1]["title"] if 0 < chapter_cursor <= len(chapters) else None,
            })

    return {"paragraphs": paragraphs, "chapters": chapters}

def main():
    src = Path("assets/booklet.docx")                     # adjust if needed
    out = Path("artifacts/booklet_index.json")
    out.parent.mkdir(parents=True, exist_ok=True)
    index = build_index(str(src))
    out.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out}")

if __name__ == "__main__":
    sys.exit(main())
